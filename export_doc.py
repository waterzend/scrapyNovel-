# -*- coding: utf-8 -*-
import pymysql
from docx import Document
import re
from docx.shared import Pt


def clear(html):
    pattern = re.compile('<div id="content[\s|\S]*script></div>')
    list_html = pattern.split(html)
    fileter_html = filter(None, list_html)
    pattern_other = re.compile('[<br>|</div>]')
    clear = pattern_other.split(fileter_html[0])
    return "".join(clear)


def write_doc(i):
    db = pymysql.connect("127.0.0.1", "root", "root", "yunxizhuan", charset='utf8')
    cursor = db.cursor()
    sql = "select * from content limit " + str(i*300) + ",300 "
    print(sql)
    document = Document()
    document.add_heading(u'芸汐传', 0)
    try:
        cursor.execute(sql)
        result = cursor.fetchone()
        while result != None:
            title = result[1]
            document.add_heading(title, 1)
            content = clear(result[2])
            font_name = u'宋体'
            paragraph = document.add_paragraph()
            run = paragraph.add_run(content)
            run.font.size = Pt(13)
            run.font.name = font_name
            result = cursor.fetchone()

    except:
        print("error")
    document.save('yunxi' + str(i) + '.docx')
    db.close()


for i in range(0, 5):
    write_doc(i)
